from io import BytesIO
import os
import asyncio
from docx import Document
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_openai import ChatOpenAI
import pandas as pd
import requests
import streamlit as st
from crewai import Agent, Task, Crew
from llama_index.core import VectorStoreIndex, Settings
from llama_index.core.node_parser import SentenceSplitter
from llama_parse import LlamaParse
import nest_asyncio
nest_asyncio.apply()

def verify_gemini_api_key(api_key):
    API_VERSION = 'v1'
    api_url = f"https://generativelanguage.googleapis.com/{API_VERSION}/models?key={api_key}"
    
    try:
        response = requests.get(api_url, headers={'Content-Type': 'application/json'})
        response.raise_for_status()  # Raises an HTTPError for bad responses
        
        # If we get here, it means the request was successful
        return True
    
    except requests.exceptions.HTTPError as e:
        
        return False
    
    except requests.exceptions.RequestException as e:
        # For any other request-related exceptions
        raise ValueError(f"An error occurred: {str(e)}")

def verify_gpt_api_key(api_key):
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    # Using a simple request to the models endpoint
    response = requests.get("https://api.openai.com/v1/models", headers=headers)
    
    if response.status_code == 200:
        return True
    elif response.status_code == 401:
        return False
    else:
        print(f"Unexpected status code: {response.status_code}")
        return False

def init_parser(api_key):
    return(LlamaParse(api_key=api_key,result_type="markdown", verbose=True))

def generate_text(parser, llm, file_path):
    Settings.chunk_size = 512
    
    if mod == 'Gemini':
        Settings.llm = llm
        Settings.embed_model = "local:BAAI/bge-small-en-v1.5"
    
    documents = parser.load_data(file_path)
    index = VectorStoreIndex.from_documents(documents, transformations=[SentenceSplitter(chunk_size=512)])
    query_engine = index.as_query_engine()
    result = query_engine.query("Could you very concisely summarize the given content? Return your response which covers the key points of the text and does not miss anything important, please.")
    
    return result['text'] if isinstance(result, dict) and 'text' in result else str(result)

def formatter(result, llm):
    agent_formatter = Agent(
        role='Summary Formatter',
        goal='Make the summary clear, concise, and well-structured without missing out on any important details',
        backstory='An expert in distilling information into readable formats.',
        verbose=True,
        allow_delegation=False,
        llm=llm
    )

    task_formatter = Task(
        description=f'''Format the given summary text to be clear and structured with headings and readable paragraphs.
                       The given summary is: {result}''',
        agent=agent_formatter,
        expected_output='''A well-structured summary with clear headings and organized content.
                           Nothing important should be missed out on from the summary'''
    )
    
    crew = Crew(
            agents=[agent_formatter],
            tasks=[task_formatter],
            verbose=2
        )
    
    return crew.kickoff()

async def process_file(file, parser, llm, upload_directory):
    temp_file_path = os.path.join(upload_directory, file.name)
    with open(temp_file_path, "wb") as f:
        f.write(file.getbuffer())
    
    generated_content = generate_text(parser, llm, temp_file_path)
    formatted_content = formatter(generated_content, llm)
    
    os.remove(temp_file_path)
    
    return {
        'File Name': file.name,
        'Summary': formatted_content
    }

async def process_all_files(files, parser, llm, upload_directory):
    tasks = [process_file(file, parser, llm, upload_directory) for file in files]
    return await asyncio.gather(*tasks)

def main():
    global mod
    mod = None
    validity_model = False
     
    st.header('Summary Generator') 
    
    # Initialize session state for storing summaries
    if 'summaries' not in st.session_state:
        st.session_state.summaries = []
          
    with st.sidebar:
        with st.form('OpenAI,Gemini'):
            model = st.radio('Choose Your LLM', ('Gemini','OpenAI'))
            api_key = st.text_input(f'Enter your API key', type="password")
            llamaindex_api_key = st.text_input(f'Enter your llamaParse API key', type="password")
            submitted = st.form_submit_button("Submit")

        if api_key and llamaindex_api_key:
            if model == "Gemini":
                validity_model = verify_gemini_api_key(api_key)
                if validity_model ==True:
                    st.write(f"Valid {model} API key")
                else:
                    st.write(f"Invalid {model} API key")
            elif model == "OpenAI":
                validity_model = verify_gpt_api_key(api_key)
                if validity_model ==True:
                    st.write(f"Valid {model} API key")
                else:
                    st.write(f"Invalid {model} API key")   
                                     
    if validity_model and llamaindex_api_key:
        if model == 'OpenAI':
            async def setup_OpenAI():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                os.environ["OPENAI_API_KEY"] = api_key
                llm = ChatOpenAI(model='gpt-4-turbo',temperature=0.6, max_tokens=2000,api_key=api_key)
                print("OpenAI Configured")
                return llm

            llm = asyncio.run(setup_OpenAI())
            mod = 'OpenAI'

        elif model == 'Gemini':
            async def setup_gemini():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                llm = ChatGoogleGenerativeAI(
                    model="gemini-1.5-flash",
                    verbose=True,
                    temperature=0.6,
                    google_api_key=api_key
                )
                print("Gemini Configured")
                return llm

            llm = asyncio.run(setup_gemini())
            mod = 'Gemini'

        parser = init_parser(llamaindex_api_key)
        
        upload_directory = "Saved Files"
        if not os.path.exists(upload_directory):
            os.makedirs(upload_directory)

        uploaded_files = st.file_uploader("Choose files", type=None, accept_multiple_files=True)

        if uploaded_files and st.button("Process Files"):
            if not st.session_state.summaries:  
                with st.spinner("Generating summaries for all files..."):
                    st.session_state.summaries = asyncio.run(process_all_files(uploaded_files, parser, llm, upload_directory))
            
            if st.session_state.summaries:
                # Display summaries in a table
                df = pd.DataFrame(st.session_state.summaries)
                st.table(df)
                
                doc = Document()
                for summary in st.session_state.summaries:
                    doc.add_heading(f'Summary for {summary["File Name"]}', level=1)
                    doc.add_paragraph(summary['Summary'])
                    doc.add_page_break()
                
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.download_button(
                    label="Download All Summaries as Word Document",
                    data=buffer,
                    file_name="All_Summaries.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
